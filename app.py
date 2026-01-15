import streamlit as st
import docx
from docx.shared import Pt, Inches, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import random
import io
import re

# ==========================================
# 頁面與常數設定
# ==========================================
st.set_page_config(
    page_title="物理題庫自動組卷系統", 
    layout="wide", 
    page_icon="🧲"
)

SOURCES = ["一般試題", "學測題", "分科測驗", "北模", "全模", "中模"]

PHYSICS_CHAPTERS = {
    "第一章.科學的態度與方法": [
        "1-1 科學的態度", "1-2 科學的方法", "1-3 國際單位制", "1-4 物理學簡介"
    ],
    "第二章.物體的運動": [
        "2-1 物體的運動", "2-2 牛頓三大運動定律", "2-3 生活中常見的力", "2-4 天體運動"
    ],
    "第三章. 物質的組成與交互作用": [
        "3-1 物質的組成", "3-2 原子的結構", "3-3 基本交互作用"
    ],
    "第四章.電與磁的統一": [
        "4-1 電流磁效應", "4-2 電磁感應", "4-3 電與磁的整合", "4-4 光波的特性", "4-5 都卜勒效應"
    ],
    "第五章. 能　量": [
        "5-1 能量的形式", "5-2 微觀尺度下的能量", "5-3 能量守恆", "5-4 質能互換"
    ],
    "第六章.量子現象": [
        "6-1 量子論的誕生", "6-2 光的粒子性", "6-3 物質的波動性", "6-4 波粒二象性", "6-5 原子光譜"
    ]
}

# ==========================================
# 核心邏輯類別與函式
# ==========================================

class Question:
    def __init__(self, q_type, content, options=None, answer=None, original_id=0, image_data=None, 
                 source="一般試題", chapter="", unit=""):
        self.id = original_id
        self.type = q_type  # 'Single', 'Multi', 'Fill'
        self.source = source
        self.chapter = chapter
        self.unit = unit
        self.content = content
        self.options = options if options else []
        self.answer = answer
        self.image_data = image_data

def extract_images_from_paragraph(paragraph, doc_part):
    """從 Word 段落中擷取圖片"""
    images = []
    nsmap = {
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    }
    # 嘗試尋找 blip 元素 (圖片參照)
    try:
        blips = paragraph._element.findall('.//a:blip', namespaces=nsmap)
        for blip in blips:
            embed_attr = blip.get(f"{{{nsmap['r']}}}embed")
            if embed_attr and embed_attr in doc_part.rels:
                part = doc_part.rels[embed_attr].target_part
                if "image" in part.content_type:
                    images.append(part.blob)
    except Exception as e:
        # 容錯處理
        print(f"Image extraction warning: {e}")
    return images

def parse_docx(file_bytes):
    """解析 Word 檔案 (支援 Source, Chapter, Unit 標籤，增強同一行標籤解析)"""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        doc_part = doc.part
    except Exception as e:
        st.error(f"檔案讀取失敗，請確認是否為有效的 Word 檔 (.docx)。錯誤: {e}")
        return []
    
    questions = []
    current_q = None
    state = None
    opt_pattern = re.compile(r'^\s*\(?[A-Ea-e]\)?\s*[.、]?\s*')
    q_id_counter = 1

    # 預設狀態 (會延續到下一題)
    curr_src = "一般試題"
    curr_chap = ""
    curr_unit = ""

    for para in doc.paragraphs:
        text = para.text.strip()
        found_images = extract_images_from_paragraph(para, doc_part)
        
        # 0. 偵測分類標籤 (Src, Chap, Unit)
        if text.startswith('[Src:'):
            curr_src = text.split(':')[1].replace(']', '').strip()
            continue
        if text.startswith('[Chap:'):
            curr_chap = text.split(':')[1].replace(']', '').strip()
            continue
        if text.startswith('[Unit:'):
            curr_unit = text.split(':')[1].replace(']', '').strip()
            continue
        if text.startswith('[Cat:'): # 相容舊版
            curr_unit = text.split(':')[1].replace(']', '').strip()
            continue

        # 1. 偵測新題目 [Type:...]
        if text.startswith('[Type:'):
            if current_q: questions.append(current_q)
            
            # 解析 Type
            end_idx = text.find(']')
            if end_idx != -1:
                q_type_str = text[6:end_idx].strip()
                remaining_text = text[end_idx+1:].strip()
            else:
                q_type_str = "Single"
                remaining_text = ""

            current_q = Question(
                q_type=q_type_str, 
                content="", 
                options=[], 
                answer="", 
                original_id=q_id_counter, 
                source=curr_src,
                chapter=curr_chap,
                unit=curr_unit
            )
            q_id_counter += 1
            state = None
            
            if remaining_text:
                text = remaining_text
            else:
                continue

        # 2. 狀態切換與單行內容處理
        if text.startswith('[Q]'):
            state = 'Q'
            text = text[3:].strip()
            if not text: continue

        elif text.startswith('[Opt]'):
            state = 'Opt'
            text = text[5:].strip()
            if not text: continue
            
        elif text.startswith('[Ans]'):
            state = 'Ans'
            remain_text = text.replace('[Ans]', '').strip()
            if remain_text and current_q: 
                current_q.answer = remain_text
            continue

        # 3. 填入內容
        if current_q:
            if found_images and state == 'Q':
                current_q.image_data = found_images[0]

            if not text: continue

            if state == 'Q': 
                current_q.content += text + "\n"
            elif state == 'Opt':
                clean_opt = opt_pattern.sub('', text)
                current_q.options.append(clean_opt)
            elif state == 'Ans': 
                current_q.answer += text

    if current_q: questions.append(current_q)
    return questions

def shuffle_options_and_update_answer(question):
    """打亂選項並修正答案"""
    if question.type == 'Fill': return question

    original_opts = question.options
    original_ans = question.answer.strip().upper()
    char_to_idx = {chr(65+i): i for i in range(len(original_opts))}
    
    correct_indices = []
    for char in original_ans:
        if char in char_to_idx: correct_indices.append(char_to_idx[char])
            
    correct_contents = [original_opts[i] for i in correct_indices]
    
    shuffled_opts_data = list(enumerate(original_opts))
    random.shuffle(shuffled_opts_data)
    new_options = [data[1] for data in shuffled_opts_data]
    
    new_ans_chars = []
    for content in correct_contents:
        try:
            new_idx = new_options.index(content)
            new_ans_chars.append(chr(65 + new_idx))
        except ValueError: pass
            
    new_ans_chars.sort()
    new_answer_str = "".join(new_ans_chars)

    return Question(
        question.type, question.content, new_options, new_answer_str, 
        question.id, question.image_data, 
        question.source, question.chapter, question.unit
    )

def set_font(doc, font_name='Times New Roman', east_asia_font='DFKai-SB'):
    """設定整份文件的預設字型"""
    style = doc.styles['Normal']
    style.font.name = font_name
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), east_asia_font)

def generate_word_files(selected_questions, shuffle=True, title="高中物理科 段考題"):
    """生成 Word 試卷 (優化排版)"""
    exam_doc = docx.Document()
    ans_doc = docx.Document()
    
    set_font(exam_doc)
    set_font(ans_doc)
    
    # === 試題卷檔頭設計 ===
    title_p = exam_doc.add_heading(title, 0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    table = exam_doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    
    cells = table.rows[0].cells
    cells[0].text = "班級："
    cells[1].text = "__________"
    cells[2].text = "姓名："
    cells[3].text = "__________"
    
    exam_doc.add_paragraph("")
    
    # === 答案卷檔頭 ===
    ans_doc.add_heading(f'{title} - 詳解卷', 0)
    ans_doc.add_paragraph('此卷包含答案與詳細分類資訊。\n')

    # === 題目內容 ===
    for idx, q in enumerate(selected_questions, 1):
        processed_q = q
        if shuffle and q.type in ['Single', 'Multi']:
            processed_q = shuffle_options_and_update_answer(q)
        
        # --- 試題卷 ---
        p = exam_doc.add_paragraph()
        q_type_text = {'Single': '單選', 'Multi': '多選', 'Fill': '填充'}.get(q.type, '未知')
        
        runner = p.add_run(f"{idx}. ({q_type_text}) {processed_q.content.strip()}")
        runner.bold = True
        
        if processed_q.image_data:
            try:
                img_stream = io.BytesIO(processed_q.image_data)
                exam_doc.add_picture(img_stream, width=Inches(3.5))
            except Exception as e:
                print(f"Error adding picture: {e}")

        if q.type != 'Fill':
            for i, opt in enumerate(processed_q.options):
                exam_doc.add_paragraph(f"({chr(65+i)}) {opt}")
        else:
            exam_doc.add_paragraph("______________________")
        
        exam_doc.add_paragraph("") 
        
        # --- 答案卷 ---
        ans_p = ans_doc.add_paragraph()
        ans_p.add_run(f"{idx}. ").bold = True
        ans_p.add_run(f"{processed_q.answer}")
        
        meta_info = []
        if processed_q.source and processed_q.source != "一般試題": meta_info.append(processed_q.source)
        if processed_q.unit: meta_info.append(processed_q.unit)
        elif processed_q.chapter: meta_info.append(processed_q.chapter)
            
        if meta_info:
            ans_p.add_run(f"  [{' / '.join(meta_info)}]").italic = True

    exam_io = io.BytesIO()
    ans_io = io.BytesIO()
    exam_doc.save(exam_io)
    ans_doc.save(ans_io)
    exam_io.seek(0)
    ans_io.seek(0)
    return exam_io, ans_io

# ==========================================
# Session State
# ==========================================
if 'question_pool' not in st.session_state:
    st.session_state['question_pool'] = []

# ==========================================
# Streamlit 介面
# ==========================================

st.title("🧲 物理題庫自動組卷系統 v3.3")
st.markdown("高中物理老師專用助理 | 支援 **LaTeX 公式**、**排版優化** 與 **線上題目編輯**。")

# --- 側邊欄 ---
with st.sidebar:
    st.header("📦 題庫管理")
    count = len(st.session_state['question_pool'])
    st.metric("目前題庫總數", f"{count} 題")
    
    if count > 0:
        if st.button("🗑️ 清空所有題目", type="primary"):
            st.session_state['question_pool'] = []
            st.rerun()
    
    st.divider()
    st.markdown("""
    **Word 匯入標籤說明：**
    - `[Src:學測題]` 來源
    - `[Chap:第一章...]` 章節
    - `[Unit:1-1...]` 單元
    - `[Type:Single]` 題型 (Single/Multi/Fill)
    - `[Q]` 題目開始
    - `[Opt]` 選項區域
    - `[Ans] A` 答案
    """)
    
    if st.button("📥 下載 Word 匯入範本"):
        sample_doc = docx.Document()
        sample_doc.add_paragraph("[Src:北模]")
        sample_doc.add_paragraph("[Chap:第四章.電與磁的統一]")
        sample_doc.add_paragraph("[Unit:4-1 電流磁效應]")
        sample_doc.add_paragraph("[Type:Single]\n[Q]\n(範例) 設載流導線電流為 $I$，距離導線 $r$ 處的磁場強度 $B$ 為何？\n[Opt]\n(A) 正比於 r\n(B) 反比於 r\n[Ans] B")
        sample_io = io.BytesIO()
        sample_doc.save(sample_io)
        sample_io.seek(0)
        st.download_button("點此下載 .docx 範本", sample_io, "template_v3.docx")

# --- 主畫面 ---
tab1, tab2, tab3 = st.tabs(["✍️ 手動新增題目", "📁 從 Word 匯入", "🚀 選題與匯出"])

# === Tab 1: 手動輸入 ===
with tab1:
    st.subheader("新增單一題目")
    st.caption("提示：在題目內容中使用 `$F=ma$` 語法可顯示數學公式。")
    
    col_cat1, col_cat2, col_cat3 = st.columns(3)
    with col_cat1:
        new_q_source = st.selectbox("來源", SOURCES)
    with col_cat2:
        chap_list = list(PHYSICS_CHAPTERS.keys())
        new_q_chap = st.selectbox("章節", chap_list)
    with col_cat3:
        unit_list = PHYSICS_CHAPTERS[new_q_chap]
        new_q_unit = st.selectbox("單元", unit_list)

    c1, c2 = st.columns([1, 3])
    with c1:
        new_q_type = st.selectbox("題型", ["Single", "Multi", "Fill"], format_func=lambda x: {'Single':'單選題', 'Multi':'多選題', 'Fill':'填充題'}[x])
    with c2:
        new_q_ans = st.text_input("正確答案", placeholder="選擇題填代號(如 A)，填充題填文字")

    new_q_content = st.text_area("題目內容 (支援 LaTeX)", height=100, placeholder="例如：求物體受力 $F = G \frac{Mm}{r^2}$ 的大小...")
    
    if "$" in new_q_content:
        st.markdown("**預覽效果：**")
        st.markdown(new_q_content)
    
    new_q_image = st.file_uploader("上傳圖片 (選用)", type=['png', 'jpg', 'jpeg'])
    
    new_q_options = []
    if new_q_type in ["Single", "Multi"]:
        opts_text = st.text_area("選項 (每一行一個選項)", height=150, placeholder="選項 A\n選項 B\n選項 C\n選項 D")
        if opts_text:
            new_q_options = [line.strip() for line in opts_text.split('\n') if line.strip()]

    if st.button("➕ 加入題庫", type="secondary"):
        if not new_q_content:
            st.error("請輸入題目內容")
        elif new_q_type != 'Fill' and not new_q_options:
            st.error("選擇題必須提供選項")
        else:
            q_id = len(st.session_state['question_pool']) + 1
            img_bytes = new_q_image.getvalue() if new_q_image else None

            new_q = Question(
                new_q_type, new_q_content, new_q_options, new_q_ans, q_id, 
                image_data=img_bytes, 
                source=new_q_source, 
                chapter=new_q_chap, 
                unit=new_q_unit
            )
            st.session_state['question_pool'].append(new_q)
            st.success(f"已加入題目！分類：{new_q_source} / {new_q_unit}")

# === Tab 2: Word 匯入 ===
with tab2:
    st.subheader("批次匯入題目")
    st.info("支援標籤：`[Src:來源]`, `[Chap:章節]`, `[Unit:單元]`。")
    uploaded_file = st.file_uploader("上傳 Word (.docx) 檔案", type=['docx'])
    
    if uploaded_file:
        if st.button("解析並加入題庫"):
            try:
                imported_qs = parse_docx(uploaded_file.read())
                if imported_qs:
                    st.session_state['question_pool'].extend(imported_qs)
                    st.success(f"成功匯入 {len(imported_qs)} 題！")
                else:
                    st.warning("未偵測到題目，請檢查 Word 檔內的標籤格式。")
            except Exception as e:
                st.error(f"解析失敗：{e}")

# === Tab 3: 選題與匯出 ===
with tab3:
    st.subheader("預覽與組卷")
    
    if not st.session_state['question_pool']:
        st.info("目前題庫是空的。請先從 Tab 1 新增或 Tab 2 匯入題目。")
    else:
        # 過濾器區域
        st.markdown("### 🔍 篩選題目")
        f_col1, f_col2 = st.columns(2)
        with f_col1:
            filter_chap = st.multiselect("篩選章節", list(PHYSICS_CHAPTERS.keys()))
        with f_col2:
            filter_src = st.multiselect("篩選來源", SOURCES)

        display_pool = []
        for i, q in enumerate(st.session_state['question_pool']):
            chap_match = (not filter_chap) or (q.chapter in filter_chap)
            src_match = (not filter_src) or (q.source in filter_src)
            
            if chap_match and src_match:
                display_pool.append((i, q))

        st.write(f"符合條件：{len(display_pool)} / 總題數：{len(st.session_state['question_pool'])}")

        col_ctrl, _ = st.columns([2, 8])
        with col_ctrl:
            select_all = st.checkbox("全選符合條件的題目", value=True)
        
        selected_final_indices = []
        
        st.write("---")
        
        # 顯示題目列表
        for original_idx, q in display_pool:
            col_check, col_text = st.columns([0.5, 9.5])
            with col_check:
                is_checked = st.checkbox("選", value=select_all, key=f"sel_{original_idx}", label_visibility="collapsed")
                if is_checked:
                    selected_final_indices.append(original_idx)
            
            with col_text:
                type_badge = {'Single': '🟢單選', 'Multi': '🔵多選', 'Fill': '🟠填充'}.get(q.type, '⚪未知')
                tags = f"[{q.source}] {q.unit}"
                preview_content = q.content.strip()
                preview_title = preview_content.splitlines()[0][:20] if preview_content else "(無內容)"
                
                with st.expander(f"{original_idx+1}. {tags} | {type_badge} | {preview_title}..."):
                    # === 編輯模式切換 ===
                    is_editing = st.checkbox(f"✏️ 編輯模式", key=f"edit_{original_idx}")
                    
                    if is_editing:
                        # 顯示編輯表單
                        with st.container(border=True):
                            st.caption("編輯題目屬性")
                            # 第一列：分類標籤
                            ec1, ec2, ec3 = st.columns(3)
                            
                            # 來源
                            try:
                                src_idx = SOURCES.index(q.source)
                            except ValueError:
                                src_idx = 0
                            new_src = ec1.selectbox("來源", SOURCES, index=src_idx, key=f"e_src_{original_idx}")
                            
                            # 章節
                            chap_keys = list(PHYSICS_CHAPTERS.keys())
                            try:
                                chap_idx = chap_keys.index(q.chapter)
                            except ValueError:
                                chap_idx = 0
                            new_chap = ec2.selectbox("章節", chap_keys, index=chap_idx, key=f"e_chap_{original_idx}")
                            
                            # 單元 (隨章節連動)
                            unit_list = PHYSICS_CHAPTERS[new_chap]
                            try:
                                unit_idx = unit_list.index(q.unit)
                            except ValueError:
                                unit_idx = 0
                            new_unit = ec3.selectbox("單元", unit_list, index=unit_idx, key=f"e_unit_{original_idx}")
                            
                            # 第二列：內容與答案
                            new_content = st.text_area("題目內容 (支援 LaTeX)", value=q.content, height=150, key=f"e_content_{original_idx}")
                            
                            new_options = q.options
                            if q.type != 'Fill':
                                opts_text = "\n".join(q.options)
                                new_opts_text = st.text_area("選項 (每行一個)", value=opts_text, height=100, key=f"e_opts_{original_idx}")
                                new_options = [line.strip() for line in new_opts_text.split('\n') if line.strip()]
                                
                            new_ans = st.text_input("答案", value=q.answer, key=f"e_ans_{original_idx}")
                            
                            if st.button("💾 儲存修改", key=f"save_{original_idx}"):
                                q.source = new_src
                                q.chapter = new_chap
                                q.unit = new_unit
                                q.content = new_content
                                q.options = new_options
                                q.answer = new_ans
                                st.success("修改已儲存！請重新展開此題以查看更新後的標題。")
                                st.rerun()
                    else:
                        # 顯示預覽模式 (原內容)
                        st.caption(f"分類：{q.chapter} > {q.unit}")
                        st.markdown("**題目**：")
                        st.markdown(q.content if q.content else "*(題目內容為空)*")
                        
                        if q.image_data:
                            st.image(q.image_data, caption="題目附圖", width=300)
                        if q.options:
                            for idx, opt in enumerate(q.options):
                                st.text(f"({chr(65+idx)}) {opt}")
                        st.markdown(f"**答案**：`{q.answer}`")
                    
                    if st.button("🗑️ 刪除此題", key=f"del_{original_idx}"):
                        st.session_state['question_pool'].pop(original_idx)
                        st.rerun()

        st.divider()
        st.write(f"已勾選匯出: **{len(selected_final_indices)}** 題")
        
        col_set1, col_set2 = st.columns(2)
        with col_set1:
            exam_title_input = st.text_input("試卷標題", value="高中物理科 段考題")
        with col_set2:
            do_shuffle = st.checkbox("啟用選項亂數重排", value=True)
        
        if st.button("🚀 生成 Word 試卷", type="primary", disabled=len(selected_final_indices)==0):
            final_qs = [st.session_state['question_pool'][i] for i in selected_final_indices]
            exam_file, ans_file = generate_word_files(final_qs, shuffle=do_shuffle, title=exam_title_input)
            
            col_d1, col_d2 = st.columns(2)
            with col_d1:
                st.download_button("📄 下載試題卷 (Word)", exam_file, "物理試題卷.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            with col_d2:
                st.download_button("🔑 下載詳解卷 (Word)", ans_file, "物理詳解卷.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
